"""Build the WebGUI Trading manuals.

Single-source pipeline: each manual is authored once in Markdown and this script
emits two artifacts next to it:

    <name>.html   — styled, self-contained HTML (online viewing AND opens cleanly in Word)
    <name>.docx   — native Word document (headings, auto TOC field, tables, code)

No pandoc dependency. Markdown -> HTML via the `markdown` package; HTML -> DOCX
via a BeautifulSoup walk into python-docx.

Usage:
    python build_docs.py                 # build all manuals
    python build_docs.py user-guide      # build one (folder name under docs/manuals)
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

HERE = Path(__file__).resolve().parent

# Manual folder -> (markdown filename, doc title, doc subtitle)
MANUALS = {
    "user-guide": ("user-guide.md", "WebGUI Trading with Schwab", "User Guide"),
    "technical-reference": ("technical-reference.md", "WebGUI Trading with Schwab", "Technical Reference"),
    "api-reference": ("api-reference.md", "WebGUI Trading with Schwab", "API / Developer Reference"),
}

MD_EXTENSIONS = ["extra", "tables", "fenced_code", "sane_lists", "toc", "attr_list"]

# ---------------------------------------------------------------------------
# HTML output
# ---------------------------------------------------------------------------

HTML_CSS = """
:root {
  --ink: #1a2230; --muted: #5b6678; --line: #d8dee9; --soft: #eef2f7;
  --accent: #2e6db4; --accent-soft: #eaf2fb; --code-bg: #f4f6fa;
  --pos: #1f8a4c; --neg: #c0392b; --warn: #b9770e;
}
* { box-sizing: border-box; }
body {
  font-family: -apple-system, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
  color: var(--ink); line-height: 1.62; margin: 0; background: #fbfcfe;
  -webkit-font-smoothing: antialiased;
}
.wrap { max-width: 980px; margin: 0 auto; padding: 48px 32px 96px; }
.titleblock { border-bottom: 3px solid var(--accent); padding-bottom: 18px; margin-bottom: 8px; }
.titleblock h1 { font-size: 2.1rem; margin: 0 0 4px; letter-spacing: -0.01em; }
.titleblock .sub { font-size: 1.2rem; color: var(--accent); font-weight: 600; }
.titleblock .meta { color: var(--muted); font-size: 0.85rem; margin-top: 8px; }
h1 { font-size: 1.7rem; margin: 2.4em 0 0.6em; padding-top: 0.3em; border-top: 1px solid var(--line); }
h2 { font-size: 1.32rem; margin: 1.8em 0 0.5em; color: #16243a; }
h3 { font-size: 1.08rem; margin: 1.4em 0 0.4em; color: #243450; }
h4 { font-size: 0.98rem; margin: 1.2em 0 0.3em; color: var(--muted); text-transform: uppercase; letter-spacing: 0.04em; }
p { margin: 0.6em 0; }
a { color: var(--accent); text-decoration: none; }
a:hover { text-decoration: underline; }
ul, ol { margin: 0.5em 0 0.9em; padding-left: 1.5em; }
li { margin: 0.25em 0; }
code { font-family: "SF Mono", Consolas, "Roboto Mono", monospace; font-size: 0.86em;
       background: var(--code-bg); padding: 1px 5px; border-radius: 4px; color: #344; }
pre { background: var(--code-bg); border: 1px solid var(--line); border-radius: 8px;
      padding: 14px 16px; overflow-x: auto; font-size: 0.85rem; line-height: 1.5; }
pre code { background: none; padding: 0; }
table { border-collapse: collapse; width: 100%; margin: 1em 0; font-size: 0.9rem; }
th, td { border: 1px solid var(--line); padding: 7px 11px; text-align: left; vertical-align: top; }
th { background: var(--accent-soft); font-weight: 600; }
tr:nth-child(even) td { background: var(--soft); }
blockquote { margin: 1em 0; padding: 6px 16px; border-left: 4px solid var(--accent);
             background: var(--accent-soft); color: #2a3a52; border-radius: 0 6px 6px 0; }
blockquote p { margin: 0.3em 0; }
hr { border: none; border-top: 1px solid var(--line); margin: 2em 0; }
.toc { background: #fff; border: 1px solid var(--line); border-radius: 10px;
       padding: 12px 22px; margin: 22px 0 8px; }
.toc > ul { margin: 0.3em 0; }
.toc ul { list-style: none; padding-left: 1.1em; }
.toc > ul { padding-left: 0.4em; }
.toc a { color: var(--ink); }
@media print {
  body { background: #fff; }
  .wrap { max-width: none; padding: 0; }
  h1 { page-break-before: auto; }
  pre, table, blockquote { page-break-inside: avoid; }
}
"""

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title} — {subtitle}</title>
<style>{css}</style>
</head>
<body>
<div class="wrap">
<div class="titleblock">
  <h1>{title}</h1>
  <div class="sub">{subtitle}</div>
  <div class="meta">{meta}</div>
</div>
{body}
</div>
</body>
</html>
"""


def build_html(md_text: str, title: str, subtitle: str, meta: str, out_path: Path) -> str:
    md = markdown.Markdown(extensions=MD_EXTENSIONS, extension_configs={
        "toc": {"title": "Contents", "permalink": False},
    })
    body = md.convert(md_text)
    # markdown's [TOC] marker -> a .toc div is produced by the toc extension's
    # placeholder; we rely on a literal "[TOC]" in source being replaced.
    html = HTML_TEMPLATE.format(title=title, subtitle=subtitle, meta=meta, css=HTML_CSS, body=body)
    out_path.write_text(html, encoding="utf-8")
    return body


# ---------------------------------------------------------------------------
# DOCX output
# ---------------------------------------------------------------------------

MONO = "Consolas"


def _set_cell_bg(cell, hex_color: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcpr.append(shd)


def _set_table_borders(table, color: str = "D8DEE9") -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _add_toc_field(doc: Document) -> None:
    """Insert a real Word TOC field (updates on open / F9)."""
    para = doc.add_paragraph()
    run = para.add_run()
    fldBegin = OxmlElement("w:fldChar"); fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fldSep = OxmlElement("w:fldChar"); fldSep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose “Update Field” to build the table of contents."
    fldEnd = OxmlElement("w:fldChar"); fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldSep)
    run2 = para.add_run(); run2._r.append(placeholder)
    run3 = para.add_run(); run3._r.append(fldEnd)


def _add_page_numbers(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(instr); run._r.append(f2)


def _inline_runs(paragraph, node, bold=False, italic=False, mono=False) -> None:
    """Recursively add inline content of an HTML node to a docx paragraph."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                run.bold = bold or None
                run.italic = italic or None
                if mono:
                    run.font.name = MONO
                    run.font.size = Pt(9.5)
        elif isinstance(child, Tag):
            if child.name in ("strong", "b"):
                _inline_runs(paragraph, child, bold=True, italic=italic, mono=mono)
            elif child.name in ("em", "i"):
                _inline_runs(paragraph, child, bold=bold, italic=True, mono=mono)
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = MONO
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x33, 0x44, 0x55)
            elif child.name == "a":
                run = paragraph.add_run(child.get_text())
                run.bold = bold or None
                run.italic = italic or None
                run.font.color.rgb = RGBColor(0x2E, 0x6D, 0xB4)
            elif child.name == "br":
                paragraph.add_run().add_break()
            else:
                _inline_runs(paragraph, child, bold=bold, italic=italic, mono=mono)


def _add_list(doc, ol_or_ul, level=0):
    ordered = ol_or_ul.name == "ol"
    style = "List Number" if ordered else "List Bullet"
    if level > 0:
        style = f"{style} {min(level + 1, 3)}"
    for li in ol_or_ul.find_all("li", recursive=False):
        p = doc.add_paragraph(style=style)
        # inline content excluding nested lists
        for child in li.children:
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                continue
            if isinstance(child, NavigableString):
                txt = str(child)
                if txt.strip():
                    p.add_run(txt)
            elif isinstance(child, Tag):
                tmp = BeautifulSoup(f"<span>{child}</span>", "html.parser").span
                _inline_runs(p, tmp)
        for nested in li.find_all(["ul", "ol"], recursive=False):
            try:
                _add_list(doc, nested, level + 1)
            except Exception:
                _add_list(doc, nested, level)


def _add_code_block(doc, pre):
    text = pre.get_text()
    text = text.rstrip("\n")
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.15)
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    # light shading via paragraph border-ish: use a single run with mono font
    for i, line in enumerate(text.split("\n")):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line if line else " ")
        run.font.name = MONO
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x22, 0x33, 0x44)
    # shade the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F4F6FA")
    pPr.append(shd)


def _add_table(doc, table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    ncols = max(len(r.find_all(["td", "th"])) for r in rows)
    tbl = doc.add_table(rows=0, cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = True
    for r_i, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        row = tbl.add_row()
        is_header = (tr.find("th") is not None)
        for c_i in range(ncols):
            cell = row.cells[c_i]
            cell.paragraphs[0].text = ""
            if c_i < len(cells):
                src = cells[c_i]
                tmp = BeautifulSoup(f"<span>{src.decode_contents()}</span>", "html.parser").span
                _inline_runs(cell.paragraphs[0], tmp)
                if is_header:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
            if is_header:
                _set_cell_bg(cell, "EAF2FB")
            elif r_i % 2 == 0:
                _set_cell_bg(cell, "F4F6FA")
    _set_table_borders(tbl)
    doc.add_paragraph()


def _add_blockquote(doc, bq):
    for p_tag in bq.find_all("p", recursive=False) or [bq]:
        p = doc.add_paragraph(style="Intense Quote") if "Intense Quote" in [s.name for s in doc.styles] else doc.add_paragraph()
        _inline_runs(p, p_tag)


def _heading_level(name: str) -> int:
    return {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}.get(name, 1)


def build_docx(body_html: str, title: str, subtitle: str, meta: str, out_path: Path) -> None:
    soup = BeautifulSoup(body_html, "html.parser")
    doc = Document()

    # base styles
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for lvl, sz in ((1, 18), (2, 14), (3, 12), (4, 11)):
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.color.rgb = RGBColor(0x16, 0x24, 0x3A) if lvl <= 2 else RGBColor(0x24, 0x34, 0x50)

    # title block
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = t.add_run(title); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x16, 0x24, 0x3A)
    s = doc.add_paragraph(); rs = s.add_run(subtitle); rs.font.size = Pt(15); rs.bold = True
    rs.font.color.rgb = RGBColor(0x2E, 0x6D, 0xB4)
    m = doc.add_paragraph(); rm = m.add_run(meta); rm.font.size = Pt(9); rm.italic = True
    rm.font.color.rgb = RGBColor(0x5B, 0x66, 0x78)

    doc.add_paragraph()
    h = doc.add_paragraph(); hr = h.add_run("Table of Contents"); hr.bold = True; hr.font.size = Pt(13)
    _add_toc_field(doc)
    doc.add_page_break()

    # walk top-level body nodes
    for node in soup.children:
        if isinstance(node, NavigableString):
            continue
        if not isinstance(node, Tag):
            continue
        name = node.name
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            # skip the markdown TOC's own heading if present inside .toc
            lvl = _heading_level(name)
            p = doc.add_heading(level=lvl)
            _inline_runs(p, node)
        elif name == "p":
            p = doc.add_paragraph()
            _inline_runs(p, node)
        elif name in ("ul", "ol"):
            _add_list(doc, node)
        elif name == "pre":
            _add_code_block(doc, node)
        elif name == "table":
            _add_table(doc, node)
        elif name == "blockquote":
            _add_blockquote(doc, node)
        elif name == "hr":
            doc.add_paragraph().add_run().add_break()
        elif name == "div" and "toc" in (node.get("class") or []):
            continue  # we use the native Word TOC field instead
        else:
            p = doc.add_paragraph()
            _inline_runs(p, node)

    _add_page_numbers(doc)
    doc.save(str(out_path))


# ---------------------------------------------------------------------------

def build_one(folder: str) -> None:
    md_name, title, subtitle = MANUALS[folder]
    md_path = HERE / folder / md_name
    if not md_path.exists():
        print(f"  ! skip {folder}: {md_name} not found")
        return
    md_text = md_path.read_text(encoding="utf-8")
    meta = "WebGUI Trading with Schwab · Generated 2026-06-23"
    html_body = build_html(md_text, title, subtitle, meta, md_path.with_suffix(".html"))
    build_docx(html_body, title, subtitle, meta, md_path.with_suffix(".docx"))
    print(f"  [ok] {folder}: {md_name} -> .html + .docx")


def main(argv) -> int:
    targets = argv[1:] or list(MANUALS)
    print("Building manuals:")
    for t in targets:
        if t not in MANUALS:
            print(f"  [skip] unknown manual '{t}' (choices: {', '.join(MANUALS)})")
            continue
        build_one(t)
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
